from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "doc" / "FORGE-X优化开发手册-React-TypeScript-CSharp.md"
OUTPUT = ROOT / "doc" / "FORGE-X优化开发手册-React-TypeScript-CSharp.docx"

NAVY = "16324F"
BLUE = "1677FF"
CYAN = "0B90A0"
ORANGE = "F28C28"
INK = "233142"
MUTED = "5F6B7A"
LIGHT = "EEF4F8"
LIGHT_BLUE = "EAF3FF"
BORDER = "C8D4DF"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        # CT_TcPr order: tcBorders -> shd -> noWrap/tcMar/.../vAlign.
        followers = {
            qn("w:noWrap"),
            qn("w:tcMar"),
            qn("w:textDirection"),
            qn("w:tcFitText"),
            qn("w:vAlign"),
            qn("w:hideMark"),
        }
        anchor = next((child for child in tc_pr if child.tag in followers), None)
        if anchor is None:
            tc_pr.append(shd)
        else:
            tc_pr.insert(tc_pr.index(anchor), shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        # Cell margins precede textDirection/tcFitText/vAlign/hideMark.
        followers = {
            qn("w:textDirection"),
            qn("w:tcFitText"),
            qn("w:vAlign"),
            qn("w:hideMark"),
        }
        anchor = next((child for child in tc_pr if child.tag in followers), None)
        if anchor is None:
            tc_pr.append(tc_mar)
        else:
            tc_pr.insert(tc_pr.index(anchor), tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_run_font(run, latin="Aptos", east_asia="Microsoft YaHei", size=None, color=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def add_update_fields_setting(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    heading_specs = {
        "Heading 1": (18, NAVY, 18, 8),
        "Heading 2": (14, BLUE, 14, 5),
        "Heading 3": (11.5, CYAN, 10, 3),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor.from_string(CYAN)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Cascadia Mono"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(8.4)
    code.font.color.rgb = RGBColor.from_string("1B3348")
    code.paragraph_format.left_indent = Cm(0.35)
    code.paragraph_format.right_indent = Cm(0.25)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0

    if "Compact" not in styles:
        compact = styles.add_style("Compact", WD_STYLE_TYPE.PARAGRAPH)
    else:
        compact = styles["Compact"]
    compact.base_style = normal
    compact.paragraph_format.space_after = Pt(2)
    compact.paragraph_format.line_spacing = 1.05
    compact.font.size = Pt(9.2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.65)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(17.4))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(12.5)
    table.columns[1].width = Cm(4.9)
    set_table_borders(table, color=WHITE, size="0")
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = left.add_run("FORGE·X INSIGHT  ·  优化开发手册")
    set_run_font(run, size=8.2, color=MUTED)
    run.bold = True
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right.add_run("V1.0  |  2026-08-10")
    set_run_font(run, size=8.2, color=MUTED)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 0, 0, 30, 0)
    line = header.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(0)
    p_pr = line._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("内部研发与实施依据  ·  第 ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(p, " PAGE ")
    run = p.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)


INLINE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text: str, *, base_size=None, base_color=None) -> None:
    pos = 0
    for match in INLINE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, latin="Cascadia Mono", east_asia="Microsoft YaHei", size=9, color="A23B00")
            run._element.get_or_add_rPr().append(OxmlElement("w:noProof"))
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color or NAVY)
            run.bold = True
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(f"{label}（{url}）")
            set_run_font(run, size=base_size, color=BLUE)
            run.underline = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color)


def add_callout(doc: Document, text: str, fill=LIGHT_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color=fill, size="0")
    cell = table.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), accent)
    tc_borders.append(left)
    tc_pr.append(tc_borders)
    # CT_TcPr requires borders before shading and cell margins.
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 220, 130, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, base_size=9.6, base_color=NAVY)
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for r_idx, data in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx in range(col_count):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F6F9FB")
            p = cell.paragraphs[0]
            p.style = doc.styles["Compact"]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_together = True
            text = data[c_idx] if c_idx < len(data) else ""
            add_inline(p, text, base_size=8.6, base_color=WHITE if r_idx == 0 else INK)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, code_lines: list[str], language: str) -> None:
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(4)
    label.paragraph_format.space_after = Pt(2)
    label.paragraph_format.keep_with_next = True
    run = label.add_run("架构示意（Mermaid 源码）" if language == "mermaid" else f"代码/配置示例 · {language or 'text'}")
    set_run_font(run, size=8.3, color=MUTED)
    run.bold = True
    keep_as_group = len(code_lines) <= 18
    for idx, line in enumerate(code_lines):
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = keep_as_group and idx < len(code_lines) - 1
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F1F5F8")
        p_pr.append(shd)
        run = p.add_run(line if line else " ")
        set_run_font(run, latin="Cascadia Mono", east_asia="Microsoft YaHei", size=8.1, color="153047")
        run._element.get_or_add_rPr().append(OxmlElement("w:noProof"))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(45)
    spacer.paragraph_format.space_after = Pt(0)

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(band, color=NAVY, size="0")
    set_cell_shading(band.cell(0, 0), NAVY)
    set_cell_margins(band.cell(0, 0), 130, 220, 130, 220)
    p = band.cell(0, 0).paragraphs[0]
    run = p.add_run("INDUSTRIAL SOFTWARE · MODERNIZATION PLAYBOOK")
    set_run_font(run, size=9, color=WHITE)
    run.bold = True

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(8)
    p.add_run("FORGE·X Insight\n优化开发手册")

    p = doc.add_paragraph(style="Subtitle")
    p.paragraph_format.space_after = Pt(28)
    p.add_run("React + TypeScript 前端重构与 C# 工业计算平台演进方案")

    line = doc.add_paragraph()
    p_pr = line._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "22")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    line.paragraph_format.space_after = Pt(22)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    meta.columns[0].width = Cm(3.6)
    meta.columns[1].width = Cm(12.5)
    set_table_borders(meta, color="D8E3EC", size="4")
    values = [
        ("文档版本", "V1.0"),
        ("项目基线", "FORGE·X Insight 0.19.0"),
        ("编制日期", "2026-08-10"),
        ("适用仓库", r"E:\Projects\3dprint"),
    ]
    for i, (key, value) in enumerate(values):
        set_cell_shading(meta.cell(i, 0), LIGHT)
        for j, text in enumerate((key, value)):
            cell = meta.cell(i, j)
            set_cell_margins(cell, 100, 130, 100, 130)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_run_font(run, size=9.5, color=NAVY if j == 0 else INK)
            run.bold = j == 0

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("技术路线")
    set_run_font(run, size=9, color=MUTED)
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("渐进替换  ·  契约先行  ·  双跑验证  ·  可回滚交付")
    set_run_font(run, size=11.5, color=NAVY)
    run.bold = True

    doc.add_page_break()


def add_contents(doc: Document, source_lines: list[str]) -> None:
    p = doc.add_paragraph("阅读导航", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    intro = doc.add_paragraph()
    add_inline(intro, "本页按实施顺序列出主要章节。Word 打开后可在导航窗格中按标题层级跳转。", base_color=MUTED)
    chapters = []
    for line in source_lines:
        if line.startswith("## "):
            chapters.append(line[3:].strip())
    table = doc.add_table(rows=len(chapters), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(1.25)
    table.columns[1].width = Cm(15.4)
    set_table_borders(table, color=WHITE, size="0")
    for idx, title in enumerate(chapters):
        num = re.match(r"(\d+).", title)
        number = num.group(1) if num else f"{idx + 1:02d}"
        c0, c1 = table.cell(idx, 0), table.cell(idx, 1)
        set_cell_shading(c0, NAVY if idx < 6 else BLUE)
        set_cell_shading(c1, "F4F8FB" if idx % 2 == 0 else WHITE)
        set_cell_margins(c0, 70, 80, 70, 80)
        set_cell_margins(c1, 70, 130, 70, 130)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(number.zfill(2))
        set_run_font(r0, size=8.5, color=WHITE)
        r0.bold = True
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(title)
        set_run_font(r1, size=9.2, color=NAVY)
        r1.bold = True
        prevent_row_split(table.rows[idx])
    doc.add_page_break()


def render_markdown(doc: Document, lines: list[str]) -> None:
    i = 0
    in_code = False
    code_language = ""
    code_lines: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines, code_language)
                in_code = False
                code_lines = []
                code_language = ""
            else:
                in_code = True
                code_language = line[3:].strip()
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        if line.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, parse_table(table_lines))
            continue

        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_paragraph(line[5:].strip(), style="Heading 3")
            i += 1
            continue
        if line.startswith("# ") or line.startswith("**副标题") or line.startswith("**文档版本") or line.startswith("**基线版本") or line.startswith("**编制日期") or line.startswith("**适用仓库"):
            i += 1
            continue

        if line.startswith("> "):
            quote = line[2:].strip()
            i += 1
            while i < len(lines) and lines[i].startswith("> "):
                quote += " " + lines[i][2:].strip()
                i += 1
            add_callout(doc, quote)
            continue

        checklist = re.match(r"^\s*- \[([ xX])\] (.+)$", line)
        bullet = re.match(r"^(\s*)- (.+)$", line)
        numbered = re.match(r"^(\s*)(\d+)\. (.+)$", line)
        if checklist:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.45)
            p.paragraph_format.first_line_indent = Cm(-0.35)
            p.paragraph_format.space_after = Pt(2)
            mark = "☒" if checklist.group(1).lower() == "x" else "☐"
            run = p.add_run(mark + "  ")
            set_run_font(run, size=10, color=BLUE)
            add_inline(p, checklist.group(2), base_size=9.6)
            i += 1
            continue
        if bullet:
            indent = min(len(bullet.group(1)) // 2, 3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.55 + indent * 0.5)
            p.paragraph_format.first_line_indent = Cm(-0.3)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run("•  ")
            set_run_font(run, size=10, color=BLUE)
            add_inline(p, bullet.group(2), base_size=9.6)
            i += 1
            continue
        if numbered:
            indent = min(len(numbered.group(1)) // 2, 3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6 + indent * 0.5)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(numbered.group(2) + ".  ")
            set_run_font(run, size=9.6, color=BLUE)
            run.bold = True
            add_inline(p, numbered.group(3), base_size=9.6)
            i += 1
            continue

        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[2:-2])
            set_run_font(run, size=10, color=NAVY)
            run.bold = True
            i += 1
            continue

        paragraph_lines = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip("\n")
            if not nxt.strip() or nxt.startswith(('#', '```', '|', '> ', '- ', '**')) or re.match(r"^\s*\d+\. ", nxt):
                break
            paragraph_lines.append(nxt.strip())
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))
        lookahead = i
        while lookahead < len(lines) and not lines[lookahead].strip():
            lookahead += 1
        if lookahead < len(lines) and lines[lookahead].startswith("```"):
            p.paragraph_format.keep_with_next = True


def build() -> None:
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    configure_styles(doc)
    add_update_fields_setting(doc)

    props = doc.core_properties
    props.title = "FORGE·X Insight 优化开发手册"
    props.subject = "React + TypeScript 前端重构与 C# 工业计算平台演进方案"
    props.author = "FORGE·X 项目组"
    props.keywords = "React, TypeScript, C#, ASP.NET Core, Three.js, 工业软件, 3D打印"
    props.comments = "根据 0.19.0 代码基线与审计结论编制"

    add_cover(doc)
    add_contents(doc, source_lines)
    render_markdown(doc, source_lines)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
